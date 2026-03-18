"""
Document Generator
Generates PDF, DOCX, and XLSX documents from draft text.

Documents are generated on-demand and streamed directly to user.
NO SERVER-SIDE STORAGE - privacy by design.
"""

from typing import Dict, Any, Optional, Tuple
from io import BytesIO
from datetime import datetime
from loguru import logger

# PDF Generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

# DOCX Generation
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.styles.style import _CharacterStyle

# XLSX Generation
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.worksheet.worksheet import Worksheet

from app.config import get_settings


class DocumentGenerator:
    """
    Generates documents in multiple formats.
    
    Design Principles:
    - No server-side storage
    - Documents generated in memory
    - Streamed directly to client
    - Consistent formatting across formats
    """
    
    def __init__(self):
        self.settings = get_settings()
    
    # =========================================================================
    # PDF GENERATION
    # =========================================================================
    
    def generate_pdf(
        self,
        draft_text: str,
        document_type: str,
        applicant_name: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Tuple[BytesIO, str]:
        """
        Generate PDF document.
        
        Args:
            draft_text: The complete draft text
            document_type: Type of document (rti/complaint)
            applicant_name: Applicant name for filename
            metadata: Additional metadata
        
        Returns:
            Tuple of (BytesIO buffer, filename)
        """
        logger.info(f"Generating PDF for {document_type}")
        
        buffer = BytesIO()
        
        # Create document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=1.0 * inch,
            leftMargin=1.0 * inch,
            topMargin=1.0 * inch,
            bottomMargin=1.0 * inch
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName='Times-Bold'
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=12,
            leading=18,  # 1.5 line spacing
            alignment=TA_JUSTIFY,
            fontName='Times-Roman',
            spaceAfter=12
        )
        
        subject_style = ParagraphStyle(
            'Subject',
            parent=styles['Normal'],
            fontSize=12,
            fontName='Times-Bold',
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Build content
        story = []
        
        # Process draft text line by line
        lines = draft_text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                story.append(Spacer(1, 6))
                continue
            
            # Detect subject line
            if line.lower().startswith('subject:'):
                story.append(Paragraph(line, subject_style))
            # Detect "To," header
            elif line == 'To,' or line == 'To':
                story.append(Paragraph('<b>To,</b>', body_style))
            # Regular paragraphs
            else:
                # Escape special characters for ReportLab
                safe_line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_line, body_style))
        
        # Add metadata footer
        if metadata:
            story.append(Spacer(1, 30))
            story.append(Paragraph(
                f"<i>Generated on: {metadata.get('generated_at', datetime.now().isoformat())}</i>",
                ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
            ))
        
        # Build PDF
        doc.build(story)
        
        # Reset buffer position
        buffer.seek(0)
        
        # Generate filename
        safe_name = "".join(c for c in applicant_name if c.isalnum() or c in " -_").strip()
        safe_name = safe_name.replace(" ", "_")[:30]
        date_str = datetime.now().strftime("%Y%m%d")
        filename = f"{document_type}_{safe_name}_{date_str}.pdf"
        
        logger.info(f"PDF generated successfully: {filename}")
        
        return buffer, filename
    
    # =========================================================================
    # DOCX GENERATION
    # =========================================================================
    
    def generate_docx(
        self,
        draft_text: str,
        document_type: str,
        applicant_name: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Tuple[BytesIO, str]:
        """
        Generate DOCX document.
        
        Args:
            draft_text: The complete draft text
            document_type: Type of document (rti/complaint)
            applicant_name: Applicant name for filename
            metadata: Additional metadata
        
        Returns:
            Tuple of (BytesIO buffer, filename)
        """
        logger.info(f"Generating DOCX for {document_type}")
        
        # Create document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        if hasattr(style, 'font'):
            style.font.name = 'Times New Roman'  # type: ignore[union-attr]
            style.font.size = Pt(12)  # type: ignore[union-attr]
        
        # Process draft text
        lines = draft_text.strip().split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - add spacing
                doc.add_paragraph()
                current_paragraph = None
                continue
            
            # Detect subject line
            if line.lower().startswith('subject:'):
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.bold = True
                para.paragraph_format.space_after = Pt(12)
                current_paragraph = None
            # Detect "To," header
            elif line == 'To,' or line == 'To':
                para = doc.add_paragraph()
                run = para.add_run('To,')
                run.bold = True
                current_paragraph = None
            # Regular text
            else:
                para = doc.add_paragraph(line)
                para.paragraph_format.line_spacing = 1.5
        
        # Add metadata
        if metadata:
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run(f"Generated on: {metadata.get('generated_at', datetime.now().isoformat())}")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = None  # Grey color
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Generate filename
        safe_name = "".join(c for c in applicant_name if c.isalnum() or c in " -_").strip()
        safe_name = safe_name.replace(" ", "_")[:30]
        date_str = datetime.now().strftime("%Y%m%d")
        filename = f"{document_type}_{safe_name}_{date_str}.docx"
        
        logger.info(f"DOCX generated successfully: {filename}")
        
        return buffer, filename
    
    # =========================================================================
    # XLSX GENERATION (Tracking Sheet)
    # =========================================================================
    
    def generate_xlsx(
        self,
        draft_text: str,
        document_type: str,
        applicant_name: str,
        applicant_details: Optional[Dict[str, Any]] = None,
        authority_details: Optional[Dict[str, Any]] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Tuple[BytesIO, str]:
        """
        Generate XLSX tracking sheet.
        
        This creates a structured spreadsheet for record-keeping,
        NOT the full document. Useful for tracking submissions.
        
        Args:
            draft_text: The complete draft text
            document_type: Type of document
            applicant_name: Applicant name
            applicant_details: Applicant information
            authority_details: Authority information
            metadata: Additional metadata
        
        Returns:
            Tuple of (BytesIO buffer, filename)
        """
        logger.info(f"Generating XLSX tracking sheet for {document_type}")
        
        wb = Workbook()
        ws: Worksheet = wb.active  # type: ignore[assignment]
        ws.title = "Application Tracker"
        
        # Styles
        header_font = Font(bold=True, size=12)
        title_font = Font(bold=True, size=14)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Title
        ws['A1'] = f"{document_type.upper()} APPLICATION TRACKING SHEET"
        ws['A1'].font = title_font
        ws.merge_cells('A1:D1')
        
        # Application Details Section
        row = 3
        ws[f'A{row}'] = "APPLICATION DETAILS"
        ws[f'A{row}'].font = header_font
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        details = [
            ("Application Type", document_type.replace("_", " ").title()),
            ("Date Generated", datetime.now().strftime("%d/%m/%Y")),
            ("Time Generated", datetime.now().strftime("%H:%M:%S")),
        ]
        
        for label, value in details:
            ws[f'A{row}'] = label
            ws[f'A{row}'].font = Font(bold=True)
            ws[f'B{row}'] = value
            row += 1
        
        # Applicant Details Section
        row += 1
        ws[f'A{row}'] = "APPLICANT DETAILS"
        ws[f'A{row}'].font = header_font
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        if applicant_details:
            for key, value in applicant_details.items():
                ws[f'A{row}'] = key.replace("_", " ").title()
                ws[f'A{row}'].font = Font(bold=True)
                ws[f'B{row}'] = str(value) if value else ""
                row += 1
        else:
            ws[f'A{row}'] = "Name"
            ws[f'A{row}'].font = Font(bold=True)
            ws[f'B{row}'] = applicant_name
            row += 1
        
        # Authority Details Section
        row += 1
        ws[f'A{row}'] = "AUTHORITY DETAILS"
        ws[f'A{row}'].font = header_font
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        if authority_details:
            for key, value in authority_details.items():
                ws[f'A{row}'] = key.replace("_", " ").title()
                ws[f'A{row}'].font = Font(bold=True)
                ws[f'B{row}'] = str(value) if value else ""
                row += 1
        else:
            ws[f'A{row}'] = "To be filled"
            row += 1
        
        # Tracking Section
        row += 1
        ws[f'A{row}'] = "SUBMISSION TRACKING"
        ws[f'A{row}'].font = header_font
        ws.merge_cells(f'A{row}:D{row}')
        
        row += 1
        tracking_fields = [
            ("Date Submitted", ""),
            ("Mode of Submission", "(Post/Online/In-Person)"),
            ("Reference Number", ""),
            ("Acknowledgment Received", "(Yes/No)"),
            ("Response Due Date", ""),
            ("Response Received", "(Yes/No)"),
            ("Response Date", ""),
            ("Status", "(Pending/Resolved/Appeal Required)"),
            ("Notes", ""),
        ]
        
        for label, hint in tracking_fields:
            ws[f'A{row}'] = label
            ws[f'A{row}'].font = Font(bold=True)
            ws[f'B{row}'] = hint
            ws[f'B{row}'].font = Font(italic=True, color="808080")
            row += 1
        
        # Draft Content (in a separate sheet)
        ws2 = wb.create_sheet("Draft Content")
        ws2['A1'] = "DRAFT CONTENT"
        ws2['A1'].font = title_font
        
        # Split draft into rows
        lines = draft_text.strip().split('\n')
        for i, line in enumerate(lines, start=3):
            ws2[f'A{i}'] = line
        
        # Adjust column widths
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 20
        ws.column_dimensions['D'].width = 20
        
        ws2.column_dimensions['A'].width = 100
        
        # Save to buffer
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        
        # Generate filename
        safe_name = "".join(c for c in applicant_name if c.isalnum() or c in " -_").strip()
        safe_name = safe_name.replace(" ", "_")[:30]
        date_str = datetime.now().strftime("%Y%m%d")
        filename = f"{document_type}_tracker_{safe_name}_{date_str}.xlsx"
        
        logger.info(f"XLSX generated successfully: {filename}")
        
        return buffer, filename
    
    # =========================================================================
    # UNIFIED GENERATION METHOD
    # =========================================================================
    
    def generate(
        self,
        format: str,
        draft_text: str,
        document_type: str,
        applicant_name: str,
        applicant_details: Optional[Dict[str, Any]] = None,
        authority_details: Optional[Dict[str, Any]] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Tuple[BytesIO, str, str]:
        """
        Generate document in specified format.
        
        Args:
            format: Output format (pdf, docx, xlsx)
            draft_text: The complete draft text
            document_type: Type of document
            applicant_name: Applicant name
            applicant_details: Applicant information
            authority_details: Authority information
            metadata: Additional metadata
        
        Returns:
            Tuple of (BytesIO buffer, filename, content_type)
        """
        format = format.lower().strip()
        
        if format == "pdf":
            buffer, filename = self.generate_pdf(draft_text, document_type, applicant_name, metadata)
            content_type = "application/pdf"
        
        elif format == "docx":
            buffer, filename = self.generate_docx(draft_text, document_type, applicant_name, metadata)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        elif format == "xlsx":
            buffer, filename = self.generate_xlsx(
                draft_text, document_type, applicant_name,
                applicant_details, authority_details, metadata
            )
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        
        else:
            raise ValueError(f"Unsupported format: {format}. Supported: pdf, docx, xlsx")
        
        return buffer, filename, content_type


# Singleton instance
_generator = None


def get_document_generator() -> DocumentGenerator:
    """Get singleton document generator instance"""
    global _generator
    if _generator is None:
        _generator = DocumentGenerator()
    return _generator
