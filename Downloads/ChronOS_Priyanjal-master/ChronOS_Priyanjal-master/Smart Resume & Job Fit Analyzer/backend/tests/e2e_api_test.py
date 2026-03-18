"""
End-to-End API Test for Smart Resume Analyzer
Tests the complete flow: Upload Resume -> Parse JD -> Evaluate -> Get Results
"""
import requests
import json
import os
from docx import Document
from io import BytesIO

BASE_URL = "http://localhost:8000"

def create_test_resume_docx():
    """Create a sample resume DOCX file in memory."""
    doc = Document()
    
    # Header
    doc.add_heading("John Doe", 0)
    doc.add_paragraph("Software Engineer | john.doe@email.com | (555) 123-4567")
    
    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Bachelor of Science in Computer Science")
    doc.add_paragraph("MIT - 2018-2022 | GPA: 3.8")
    
    # Experience
    doc.add_heading("Professional Experience", level=1)
    doc.add_paragraph("Software Engineer at Google (2022-Present)")
    doc.add_paragraph("- Developed microservices using Python and FastAPI")
    doc.add_paragraph("- Led team of 3 engineers on cloud migration project")
    doc.add_paragraph("- Reduced API latency by 40% through optimization")
    
    doc.add_paragraph("Software Engineering Intern at Meta (2021)")
    doc.add_paragraph("- Built React components for the News Feed")
    doc.add_paragraph("- Implemented unit tests with Jest")
    
    # Skills
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("Languages: Python, JavaScript, TypeScript, Java, SQL")
    doc.add_paragraph("Frameworks: React, FastAPI, Django, Node.js")
    doc.add_paragraph("Cloud: AWS, GCP, Docker, Kubernetes")
    doc.add_paragraph("Tools: Git, CI/CD, Terraform")
    
    # Projects
    doc.add_heading("Projects", level=1)
    doc.add_paragraph("AI Resume Analyzer")
    doc.add_paragraph("- Built a full-stack application using React and Python")
    doc.add_paragraph("- Implemented NLP-based skill extraction")
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def test_health():
    """Test health endpoint."""
    print("\n1. Testing Health Endpoint...")
    try:
        response = requests.get(f"{BASE_URL}/health")
        if response.status_code == 200:
            print(f"   ✅ Health check passed: {response.json()}")
            return True
        else:
            print(f"   ❌ Health check failed: {response.status_code}")
            return False
    except Exception as e:
        print(f"   ❌ Connection failed: {e}")
        return False

def test_upload_resume():
    """Test resume upload."""
    print("\n2. Testing Resume Upload...")
    
    resume_buffer = create_test_resume_docx()
    files = {
        'file': ('test_resume.docx', resume_buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    }
    
    try:
        response = requests.post(f"{BASE_URL}/api/upload-resume", files=files)
        if response.status_code == 200:
            data = response.json()
            session_id = data.get('session_id')
            parsed = data.get('parsed_resume', {})
            
            print(f"   ✅ Upload successful!")
            print(f"   Session ID: {session_id}")
            print(f"   Skills found: {len(parsed.get('skills', []))}")
            print(f"   Experience entries: {len(parsed.get('experience', []))}")
            print(f"   Education entries: {len(parsed.get('education', []))}")
            
            return session_id, parsed
        else:
            print(f"   ❌ Upload failed: {response.status_code}")
            print(f"   Response: {response.text[:500]}")
            return None, None
    except Exception as e:
        print(f"   ❌ Request failed: {e}")
        return None, None

def test_analyze_jd(session_id):
    """Test job description analysis."""
    print("\n3. Testing Job Description Analysis...")
    
    job_description = """
    Senior Software Engineer - Backend
    
    Company: TechCorp Inc.
    
    Requirements:
    - 3+ years of experience with Python
    - Strong knowledge of FastAPI or Django
    - Experience with AWS or GCP cloud services
    - Proficiency in SQL and database design
    - Experience with Docker and Kubernetes
    - Strong problem-solving skills
    
    Nice to have:
    - Experience with React or frontend development
    - Knowledge of machine learning
    - Team leadership experience
    
    Responsibilities:
    - Design and implement scalable backend services
    - Lead technical discussions and code reviews
    - Mentor junior engineers
    """
    
    headers = {'Content-Type': 'application/json'}
    payload = {'job_description': job_description}
    
    try:
        response = requests.post(
            f"{BASE_URL}/api/analyze-jd?session_id={session_id}",
            headers=headers,
            json=payload
        )
        
        if response.status_code == 200:
            data = response.json()
            parsed_jd = data.get('parsed_jd', {})
            
            print(f"   ✅ JD Analysis successful!")
            print(f"   Required skills: {parsed_jd.get('required_skills', [])}")
            print(f"   Optional skills: {parsed_jd.get('optional_skills', [])}")
            
            return True, parsed_jd
        else:
            print(f"   ❌ JD Analysis failed: {response.status_code}")
            print(f"   Response: {response.text[:500]}")
            return False, None
    except Exception as e:
        print(f"   ❌ Request failed: {e}")
        return False, None

def test_evaluate(session_id):
    """Test evaluation endpoint."""
    print("\n4. Testing Evaluation...")
    
    headers = {'Content-Type': 'application/json'}
    payload = {'session_id': session_id}
    
    try:
        response = requests.post(
            f"{BASE_URL}/api/evaluate",
            headers=headers,
            json=payload
        )
        
        if response.status_code == 200:
            data = response.json()
            result = data.get('result', {})
            
            print(f"   ✅ Evaluation successful!")
            print(f"   Job Fit Score: {result.get('job_fit_score')}/100")
            print(f"   Confidence: {result.get('confidence_level')}")
            print(f"   Matched skills: {result.get('matched_count')}")
            print(f"   Partial matches: {result.get('partial_count')}")
            print(f"   Missing skills: {result.get('missing_count')}")
            print(f"   Explanation: {result.get('explanation', '')[:200]}...")
            
            return True, result
        else:
            print(f"   ❌ Evaluation failed: {response.status_code}")
            print(f"   Response: {response.text[:500]}")
            return False, None
    except Exception as e:
        print(f"   ❌ Request failed: {e}")
        return False, None

def run_all_tests():
    """Run the complete E2E test suite."""
    print("=" * 60)
    print("END-TO-END API TEST SUITE")
    print("=" * 60)
    
    # Step 1: Health check
    if not test_health():
        print("\n❌ Backend is not running. Please start the server first.")
        return False
    
    # Step 2: Upload resume
    session_id, parsed_resume = test_upload_resume()
    if not session_id:
        print("\n❌ Resume upload failed. Cannot continue.")
        return False
    
    # Step 3: Analyze JD
    jd_success, parsed_jd = test_analyze_jd(session_id)
    if not jd_success:
        print("\n❌ JD analysis failed. Cannot continue.")
        return False
    
    # Step 4: Evaluate
    eval_success, result = test_evaluate(session_id)
    if not eval_success:
        print("\n❌ Evaluation failed.")
        return False
    
    print("\n" + "=" * 60)
    print("✅ ALL TESTS PASSED!")
    print("=" * 60)
    
    # Summary
    print("\nSUMMARY:")
    print(f"  Resume Skills Found: {len(parsed_resume.get('skills', []))}")
    print(f"  Resume Experience Entries: {len(parsed_resume.get('experience', []))}")
    print(f"  Resume Education Entries: {len(parsed_resume.get('education', []))}")
    print(f"  JD Required Skills: {len(parsed_jd.get('required_skills', []))}")
    print(f"  Final Job Fit Score: {result.get('job_fit_score')}/100")
    
    return True

if __name__ == "__main__":
    run_all_tests()
