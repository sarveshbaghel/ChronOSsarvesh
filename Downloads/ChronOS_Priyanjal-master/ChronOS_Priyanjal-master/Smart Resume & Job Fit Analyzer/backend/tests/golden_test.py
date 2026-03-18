import os
import sys
import logging
from datetime import datetime

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("GoldenTest")

# Add backend directory to sys.path
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

try:
    from parsers import parse_resume
    from parsers.pdf_parser import parse_pdf
    from parsers.docx_parser import parse_docx
    from rules import evaluate
    from api.schemas import ParsedJobDescription, JDRequirement, SkillPriority
    from docx import Document
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
except ImportError as e:
    logger.error(f"Import failed: {e}")
    logger.error("Please ensure you are in the 'backend' directory and dependencies are installed.")
    sys.exit(1)

def create_sample_docx(filename):
    """Create a sample DOCX resume."""
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer')
    doc.add_paragraph('john.doe@example.com | (555) 123-4567')
    
    doc.add_heading('Experience', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Senior Software Engineer').bold = True
    p.add_run(' - TechCorp Inc.').italic = True
    doc.add_paragraph('January 2020 - Present')
    doc.add_paragraph('• Led migration to microservices architecture using Python and FastAPI')
    doc.add_paragraph('• Managed AWS infrastructure with Terraform and Docker')
    doc.add_paragraph('• Mentored junior developers in Agile best practices')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science')
    doc.add_paragraph('University of Tech, 2015-2019')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, JavaScript, React, Node.js, AWS, Docker, PostgreSQL, Redis')
    
    doc.save(filename)
    logger.info(f"Created sample DOCX: {filename}")

def create_sample_pdf(filename):
    """Create a sample PDF resume."""
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    
    # Header
    c.setFont("Helvetica-Bold", 24)
    c.drawString(50, height - 50, "Jane Smith")
    
    c.setFont("Helvetica", 12)
    c.drawString(50, height - 70, "Data Scientist")
    c.drawString(50, height - 85, "jane.smith@example.com | (555) 987-6543")
    
    # Experience
    y = height - 120
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Experience")
    y -= 25
    
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, "Lead Data Scientist - DataCorp")
    y -= 15
    c.setFont("Helvetica", 10)
    c.drawString(50, y, "March 2019 - Present")
    y -= 15
    
    c.drawString(60, y, "• Developed predictive models using Scikit-Learn and TensorFlow")
    y -= 15
    c.drawString(60, y, "• Optimized data pipelines with Spark and Airflow")
    y -= 25
    
    # Education
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Education")
    y -= 25
    c.setFont("Helvetica", 12)
    c.drawString(50, y, "Master of Science in Data Science")
    y -= 15
    c.drawString(50, y, "Stanford University, 2018")
    y -= 25
    
    # Skills
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Skills")
    y -= 25
    c.setFont("Helvetica", 10)
    c.drawString(50, y, "Python, R, SQL, TensorFlow, PyTorch, Tableau, Power BI, AWS")
    
    c.save()
    logger.info(f"Created sample PDF: {filename}")

def verify_resume(parsed_resume, expected_name, expected_skills):
    """Verify parsed resume data."""
    logger.info(f"--- Verifying {expected_name} ---")
    
    # Check basic parsing
    if not parsed_resume.raw_text:
        logger.error("❌ FAILED: Raw text is empty")
        return False
    
    logger.info(f"Raw text length: {len(parsed_resume.raw_text)} chars")
    
    # Check skills
    extracted_skills = [s.canonical_name for s in parsed_resume.skills]
    logger.info(f"Extracted Skills: {', '.join(extracted_skills)}")
    
    missing_skills = [s for s in expected_skills if s not in extracted_skills]
    
    if missing_skills:
        logger.warning(f"⚠️  MISSING SKILLS: {', '.join(missing_skills)}")
        # Check case insensitive
        extracted_lower = [s.lower() for s in extracted_skills]
        really_missing = [s for s in missing_skills if s.lower() not in extracted_lower]
        if really_missing:
            logger.error(f"❌ FAILED: Critical skills missing: {really_missing}")
            return False
    
    logger.info("✅ Skills extraction looks good")
    
    # Check sections
    if parsed_resume.experience:
        logger.info(f"✅ Experience entries found: {len(parsed_resume.experience)}")
        logger.info(f"   First Company: {parsed_resume.experience[0].company}")
        logger.info(f"   First Title: {parsed_resume.experience[0].title}")
    else:
        logger.warning("⚠️  No experience entries found")
        
    if parsed_resume.education:
        logger.info(f"✅ Education entries found: {len(parsed_resume.education)}")
        logger.info(f"   Institution: {parsed_resume.education[0].institution}")
    else:
        logger.warning("⚠️  No education entries found")
        
    return True

def verify_evaluation(parsed_resume):
    """Verify evaluation logic."""
    logger.info("--- Verifying Evaluation ---")
    
    # Create valid JD
    jd = ParsedJobDescription(
        raw_text="Sample JD",
        required_skills=["Python", "AWS"],
        optional_skills=["Docker", "React"],
        requirements=[
            JDRequirement(text="Python", skills=["Python"], priority=SkillPriority.REQUIRED),
            JDRequirement(text="AWS", skills=["AWS"], priority=SkillPriority.REQUIRED),
        ]
    )
    
    try:
        result = evaluate(parsed_resume, jd)
        logger.info(f"✅ Evaluation successful. Score: {result.job_fit_score}")
        
        if result.confidence_level:
             logger.info(f"✅ Confidence Level: {result.confidence_level}")
             
        if result.experience_signals:
            logger.info(f"✅ Experience Signals: Ownership={result.experience_signals.ownership_strength}")
            
        return True
    except Exception as e:
        logger.error(f"❌ Evaluation crashed: {e}")
        import traceback
        traceback.print_exc()
        return False

def run_golden_test():
    """Run the golden test suite."""
    print("\n🌟 RUNNING GOLDEN TEST SUITE 🌟\n")
    
    # Ensure uploads directory exists
    os.makedirs("tests/golden_temp", exist_ok=True)
    
    docx_path = "tests/golden_temp/golden_resume.docx"
    pdf_path = "tests/golden_temp/golden_resume.pdf"
    
    # 1. Test DOCX + Evaluation
    create_sample_docx(docx_path)
    try:
        logger.info("Parsing DOCX...")
        parsed_docx = parse_resume(docx_path, "docx")
        expected_docx_skills = ["Python", "FastAPI", "AWS", "Docker", "React"]
        
        parsing_ok = verify_resume(parsed_docx, "John Doe (DOCX)", expected_docx_skills)
        eval_ok = verify_evaluation(parsed_docx)
        
        if parsing_ok and eval_ok:
            print("\n✅ DOCX + EVAL TEST PASSED\n")
        else:
            print("\n❌ DOCX TEST FAILED\n")
            
    except Exception as e:
        logger.error(f"DOCX Parsing crashed: {e}")
        import traceback
        traceback.print_exc()

    # 2. Test PDF
    create_sample_pdf(pdf_path)
    try:
        logger.info("Parsing PDF...")
        parsed_pdf = parse_resume(pdf_path, "pdf")
        expected_pdf_skills = ["Python", "TensorFlow", "SQL", "AWS"]
        if verify_resume(parsed_pdf, "Jane Smith (PDF)", expected_pdf_skills):
            print("\n✅ PDF TEST PASSED\n")
        else:
            print("\n❌ PDF TEST FAILED\n")
            
    except Exception as e:
        logger.error(f"PDF Parsing crashed: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    run_golden_test()
