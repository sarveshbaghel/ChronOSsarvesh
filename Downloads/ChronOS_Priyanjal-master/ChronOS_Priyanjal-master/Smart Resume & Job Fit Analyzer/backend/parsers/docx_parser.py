"""
DOCX parsing module using python-docx.
Extracts text preserving document structure, including from text boxes and shapes.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from typing import Tuple
import re


def parse_docx(file_path: str) -> Tuple[str, list[dict]]:
    """
    Parse a DOCX file and extract text with structure.
    Handles multi-column layouts, text boxes, and shapes.
    
    Args:
        file_path: Path to the DOCX file
    
    Returns:
        Tuple of (raw_text, text_blocks)
        - raw_text: Full concatenated text
        - text_blocks: List of text blocks with metadata
    """
    doc = Document(file_path)
    text_blocks = []
    raw_text_parts = []
    line_num = 0
    
    # Method 1: Extract from paragraphs (standard content)
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        
        line_num += 1
        raw_text_parts.append(para.text)
        
        # Determine if heading or bold
        is_heading = para.style.name.lower().startswith("heading") if para.style else False
        is_bold = _is_paragraph_bold(para)
        font_size = _get_paragraph_font_size(para)
        
        text_blocks.append({
            "text": para.text.strip(),
            "page": 1,
            "line": line_num,
            "top": line_num * 20,
            "left": 0,
            "font_size": font_size,
            "is_bold": is_bold or is_heading,
            "is_heading": is_heading,
            "style": para.style.name if para.style else None,
        })
    
    # Method 2: Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    line_num += 1
                    raw_text_parts.append(cell.text)
                    text_blocks.append({
                        "text": cell.text.strip(),
                        "page": 1,
                        "line": line_num,
                        "top": line_num * 20,
                        "left": 0,
                        "font_size": None,
                        "is_bold": False,
                        "is_table_cell": True,
                    })
    
    # Method 3: Extract from text boxes, shapes, and all XML elements
    # This catches content in multi-column layouts that use shapes/text boxes
    body = doc._element.body
    all_xml_text = _extract_all_xml_text(body)
    
    # Only add if we didn't get much from paragraphs/tables
    if len(raw_text_parts) < 3 or len(' '.join(raw_text_parts)) < 100:
        # Use the XML extraction as the primary source
        for text_chunk in all_xml_text:
            if text_chunk.strip() and text_chunk.strip() not in raw_text_parts:
                line_num += 1
                raw_text_parts.append(text_chunk.strip())
                text_blocks.append({
                    "text": text_chunk.strip(),
                    "page": 1,
                    "line": line_num,
                    "top": line_num * 20,
                    "left": 0,
                    "font_size": None,
                    "is_bold": False,
                    "is_textbox": True,
                })
    
    raw_text = "\n".join(raw_text_parts)
    return raw_text, text_blocks


def _extract_all_xml_text(element) -> list[str]:
    """
    Recursively extract all text from XML elements.
    This catches text in text boxes, shapes, and other special containers.
    """
    texts = []
    
    for elem in element.iter():
        # Get text content
        if elem.text and elem.text.strip():
            texts.append(elem.text.strip())
        if elem.tail and elem.tail.strip():
            texts.append(elem.tail.strip())
    
    # Deduplicate while preserving order
    seen = set()
    unique_texts = []
    for text in texts:
        # Normalize whitespace for comparison
        normalized = ' '.join(text.split())
        if normalized not in seen and len(normalized) > 1:
            seen.add(normalized)
            unique_texts.append(normalized)
    
    return unique_texts


def _is_paragraph_bold(para) -> bool:
    """
    Check if the majority of the paragraph is bold.
    """
    if not para.runs:
        return False
    
    bold_count = sum(1 for run in para.runs if run.bold)
    return bold_count > len(para.runs) / 2


def _get_paragraph_font_size(para) -> float | None:
    """
    Get the font size of the paragraph.
    """
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
    return None

